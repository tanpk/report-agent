import os
import json
import hashlib
from google import genai
from google.genai import types
from dotenv import load_dotenv
from file_reader import read_file

load_dotenv()

# --- 設定 ---
DEBUG = False       # True にするとAPIを呼ばずダミー出力を返す
MAX_TOKENS = 10240  # 出力トークン上限（コスト削減の要）
CACHE_DIR = ".cache" # キャッシュ保存先

# --- モデル設定 ---
MODEL_HIGH = "gemini-3-flash-preview"   # 高品質生成（要点まとめ・構成案）
MODEL_LOW  = "gemini-3.1-flash-lite-preview"   # 軽量処理（OCR等）

# system_promptを短く（トークン削減）
SYSTEM_PROMPT = (
    "大学生のレポート作成支援の専門家。"
    "資料に基づいてのみ回答し、記載のない情報は「資料に記載なし」と示す。"
    "構成は「1.目的 2.原理 3.実験装置および実験方法 4.結果 5.考察 6.まとめ」の順。"
    "学術的文体。出力はWord変換用にMarkdown形式(##/###/-/**)で記述。"
)

class ReportAgent:
    def __init__(self):
        self.client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
        os.makedirs(CACHE_DIR, exist_ok=True)

    # --- キャッシュ ---

    def _cache_key(self, text: str) -> str:
        """入力テキストのSHA256ハッシュをキャッシュキーとして返す"""
        return hashlib.sha256(text.encode()).hexdigest()[:16]

    def _load_cache(self, key: str) -> str | None:
        """キャッシュが存在すれば読み込む"""
        path = os.path.join(CACHE_DIR, f"{key}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                print("  [キャッシュ使用]")
                return json.load(f)["result"]
        return None

    def _save_cache(self, key: str, result: str):
        """結果をキャッシュに保存する"""
        path = os.path.join(CACHE_DIR, f"{key}.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump({"result": result}, f, ensure_ascii=False)

    # --- API呼び出し（共通化） ---

    def _generate(self, prompt: str, max_tokens: int | None = None) -> str:
        """APIを呼び出す共通メソッド。DEBUGモード・キャッシュに対応"""
        key = self._cache_key(prompt)

        # デバッグモード：APIを呼ばずにダミーを返す
        if DEBUG:
            print("  [DEBUGモード：API呼び出しをスキップ]")
            return f"[DEBUG] プロンプト長: {len(prompt)}文字"

        # キャッシュヒット：そのまま返す
        cached = self._load_cache(key)
        if cached:
            return cached

        # max_tokensの決定（Noneなら無制限＝指定なし）
        token_limit = max_tokens if max_tokens is not None else MAX_TOKENS
        config_kwargs = {"system_instruction": SYSTEM_PROMPT}
        if token_limit:
            config_kwargs["max_output_tokens"] = token_limit

        # API呼び出し
        response = self.client.models.generate_content(
            model=MODEL_HIGH,
            contents=prompt,
            config=types.GenerateContentConfig(**config_kwargs)
        )
        result = response.text
        self._save_cache(key, result)
        return result

    # --- ファイル読み込み ---

    def load_files(self, file_paths: list[str], force_ocr: bool = False) -> str:
        """複数ファイルを読み込んで1つのテキストに結合する"""
        parts = []
        for path in file_paths:
            print(f"読み込み中: {path}")
            content = read_file(path, force_ocr=force_ocr)
            filename = os.path.basename(path)
            parts.append(f"[{filename}]\n{content}")  # ラベルを短縮
        return "\n\n".join(parts)

    # --- 生成（1回のAPI呼び出しで要約+構成案を同時生成） ---

    def summarize(self, content: str) -> str:
        """要点まとめ"""
        prompt = f"以下の資料の要点をまとめてください。\n\n{content}"
        return self._generate(prompt)

    def suggest_structure(self, content: str, theme: str) -> str:
        """レポート構成案を提案する"""
        prompt = f"テーマ：{theme}\n\n以下の資料をもとにレポートの構成案を提案してください。\n\n{content}"
        return self._generate(prompt)

    def summarize_and_structure(self, content: str, theme: str, chapter_instruction: str = "", max_tokens: int | None = None, output_summary: bool = True, output_report: bool = True) -> tuple[str, str]:
        """要約とレポート全文を生成する。output_summary/output_reportで個別に制御可能"""
        chapter_note = f"\n{chapter_instruction}" if chapter_instruction else ""

        summary = ""
        structure = ""

        # ①要点まとめ
        if output_summary:
            summary_prompt = f"以下の資料の要点をまとめてください。\n\n{content}"
            summary = self._generate(summary_prompt, max_tokens=max_tokens)

        # ②レポート全文
        if output_report:
            structure_prompt = (
                f"テーマ：{theme}\n\n"
                f"以下の資料をもとに、レポートを完全に記述してください。"
                f"各章の内容を詳細に執筆し、構成案ではなく実際のレポート本文として仕上げてください。{chapter_note}\n\n{content}"
            )
            structure = self._generate(structure_prompt, max_tokens=max_tokens)

        return summary, structure

    # --- 保存 ---

    def save_docx(self, content: str, filename: str):
        """結果をdocx形式で保存する"""
        from docx import Document
        doc = Document()
        for line in content.splitlines():
            s = line.strip()
            if not s:
                continue
            if s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s.startswith(("- ", "* ", "・")):
                doc.add_paragraph(s[2:], style="List Bullet")
            else:
                doc.add_paragraph(s.replace("**", ""))
        doc.save(filename)
        print(f"{filename}に保存しました。")
